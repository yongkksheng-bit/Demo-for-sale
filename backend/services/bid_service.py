from langchain.chat_models import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.schema.output_parser import StrOutputParser
import json
from config import settings
import PyPDF2
import docx
import io


class BidAnalysisService:
    """招标文件分析服务"""
    
    def __init__(self):
        self.llm = ChatOpenAI(
            openai_api_key=settings.ARK_API_KEY,
            openai_api_base=settings.ARK_BASE_URL,
            model_name=settings.ARK_MODEL,
            temperature=0.3,
            max_tokens=4096,
            request_timeout=120
        )
    
    def analyze_bid_document(self, file_content: bytes, filename: str, identity: str = "大客户销售") -> dict:
        """
        分析招标文件，提炼关键需求和风险点
        
        Args:
            file_content: 文件内容（字节）
            filename: 文件名
            identity: 销售身份
        
        Returns:
            分析结果字典
        """
        # 解析文件内容
        text = self._extract_text(file_content, filename)
        
        if not text or len(text.strip()) < 50:
            raise ValueError("文件内容为空或无法解析")
        
        # 限制文本长度，避免超出token限制
        if len(text) > 20000:
            text = text[:20000] + "\n...（内容过长，已截断）"
        
        # 调用大模型分析
        prompt = ChatPromptTemplate.from_messages([
            ("system", f"""你是一位资深的投标分析师和{identity}专家，擅长从招标文件中快速提炼关键信息和风险点。

请仔细阅读以下招标文件内容，提炼出关键信息，并以JSON格式返回结果。

【输出格式要求】
严格按照以下JSON格式返回，不要有其他文字：
{{
    "project_name": "项目名称",
    "buyer": "采购方/招标人",
    "budget": "预算金额（如果有）",
    "duration": "交付周期/工期",
    "deadline": "投标截止时间",
    "location": "项目地点",
    "technical_requirements": [
        "技术需求1",
        "技术需求2"
    ],
    "service_requirements": [
        "服务需求1", 
        "服务需求2"
    ],
    "business_requirements": [
        "商务要求1",
        "商务要求2"
    ],
    "risks": [
        {{
            "title": "风险点标题",
            "description": "风险描述",
            "suggestion": "应对建议"
        }}
    ],
    "suggestions": [
        "应对建议1",
        "应对建议2"
    ]
}}

【分析要点】
1. 项目基本信息：从文件中提取项目名称、采购方、预算、工期、截止时间、地点等关键信息
2. 技术需求：提炼核心技术要求、功能需求、性能指标等
3. 服务需求：提炼运维、培训、售后、实施等服务要求
4. 商务要求：提炼资质要求、付款方式、验收标准、违约责任等
5. 风险点：识别潜在风险，如：
   - 资质门槛过高或排他性条款
   - 评分标准明显偏向某家厂商
   - 技术参数有指定品牌或倾向性
   - 工期过于紧张
   - 付款条件苛刻
   - 验收标准模糊
   - 知识产权条款不利
   每个风险点都要给出具体的应对建议
6. 应对建议：给出整体的投标策略建议

【注意】
- 只基于文件内容进行分析，不要编造信息
- 如果某个信息找不到，填"待确认"
- 需求和风险点要具体、明确，不要空泛
- 站在{identity}的角度，重点关注和销售相关的风险点和机会点
- 严格输出JSON格式，不要有其他文字说明"""),
            ("user", """【招标文件内容】
{text}

请分析以上招标文件内容，按照要求输出JSON格式的分析结果。""")
        ])
        
        chain = prompt | self.llm | StrOutputParser()
        
        print(f"📄 正在分析招标文件: {filename}")
        result = chain.invoke({
            "text": text,
            "identity": identity
        })
        
        # 解析JSON结果
        try:
            # 尝试提取JSON部分（有时候大模型会加```json标记）
            if "```json" in result:
                result = result.split("```json")[1].split("```")[0].strip()
            elif "```" in result:
                result = result.split("```")[1].strip()
            
            data = json.loads(result)
            return data
        except json.JSONDecodeError as e:
            print(f"⚠️ JSON解析失败: {e}")
            print(f"原始结果: {result[:500]}")
            # 如果解析失败，返回一个基本结构
            return {
                "project_name": filename,
                "buyer": "待确认",
                "budget": "待确认",
                "duration": "待确认",
                "deadline": "待确认",
                "location": "待确认",
                "technical_requirements": ["文件解析成功，但JSON格式解析失败，请查看原始文件"],
                "service_requirements": [],
                "business_requirements": [],
                "risks": [],
                "suggestions": ["建议人工复核文件内容"]
            }
    
    def _extract_text(self, file_content: bytes, filename: str) -> str:
        """从文件中提取文本内容"""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return self._extract_pdf_text(file_content)
        elif filename_lower.endswith('.docx'):
            return self._extract_docx_text(file_content)
        elif filename_lower.endswith('.doc'):
            # .doc 格式比较复杂，暂时提示不支持
            return "暂不支持 .doc 格式，请转换为 .docx 或 PDF 后再上传"
        elif filename_lower.endswith('.txt'):
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return file_content.decode('gbk')
                except:
                    return "文件编码无法识别"
        else:
            return f"不支持的文件格式: {filename}"
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """提取PDF文本"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
        except Exception as e:
            print(f"PDF解析失败: {e}")
            return f"PDF解析失败: {str(e)}"
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """提取Word文档文本"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
            return text
        except Exception as e:
            print(f"Word解析失败: {e}")
            return f"Word文档解析失败: {str(e)}"


# 全局单例
bid_analysis_service = BidAnalysisService()
